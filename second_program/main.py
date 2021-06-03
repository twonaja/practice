from shutil import copy
import yaml
import docx
import openpyxl


# функция которая возвращает список индексов у которых в названии
# содержится сигнатура заданая пользователем data_list -> список
# с данными, sign -> сигнатура для поиска
def find_match_indexes(data_list: list, sign: str):
    length_dl = len(data_list)
    index_list = []
    for i in range(0, length_dl):
        if sign in data_list[i]:
            index_list.append(i)
    return index_list



# функция используется для замены слов и
# зачистки .docx документов - от оставшихся сигнатур
def docx_replace(doc_name: str, signature: str, repl_data: str):
    word_docx = docx.Document(doc_name)
    # поиск сигнатур в обычном тексте
    sign_list_len = len(signature)
    for paragraph in word_docx.paragraphs:
        # $1 - сигнатура на замену || const1 - вставляемое слово
        paragraph.text = paragraph.text.replace(signature, repl_data)
        # поиск сигнатур в имеющихся таблицах, актуально к примеру для титульных листов
        for table in word_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(signature, repl_data)
    word_docx.save(doc_name)


# функция используется для замены слов и
# зачистки .xlsx документов - от оставшихся сигнатур
def excel_replace(doc_name: str, signature: str, repl_data: str):
    excel_doc = openpyxl.load_workbook(doc_name)
    # Лист1 (имя листа в шаблоне не забыть найти более универсальный метод)
    ws = excel_doc["Лист1"]  # открытие листа
    # поиск по всем доступным колонкам и стобцам
    for r in range(1, ws.max_row + 1):
        for c in range(1, ws.max_column + 1):
            # получает строку и проверяет ее, если она не пустая пытается выполнить замену
            s = ws.cell(r, c).value
            if s != None and type(s) != int:
                ws.cell(r, c).value = s.replace(signature, repl_data)
    excel_doc.save(doc_name)


if __name__ == '__main__':
    tmpList = []
    print("Start program...\n\n")
    # часть кода для работы с yaml файлом
    with open("script.yaml", 'r', encoding='utf-8') as stream:
        try:
            load = yaml.safe_load(stream)
            # имена шаблонов - тип список, [0] - .docx; [1] - .xlsx.
            examples_name = load['examples_name']
            # сигнатуры
            list_of_signature = load['list_of_signature']
            # имена новых файлов и вставляемые данные - тип список
            names_and_data = load['names_and_data']
            num_of_docx = load['num_of_docx']
            num_of_xlsx = load['num_of_xlsx']
        except yaml.YAMLError as exc:
            print(exc)
        # конец работы с yaml файлом

    len_of_sign_list = len(list_of_signature)
    len_of_nd_list = len(names_and_data)

    const_for_add = (len_of_sign_list + 1)

    j = 0
    for i in range(0, num_of_docx):
        copy(examples_name[0], names_and_data[j] + '.docx')
        j += const_for_add

    j = 0
    for i in range(0, num_of_xlsx):
        copy(examples_name[1], names_and_data[j] + '.xlsx')
        j += const_for_add

    j = 0
    for i in range(0, num_of_docx):
        tmpList = names_and_data[j + 1: j + const_for_add]  # получаем временный список данных для замены
        num_of_repl_str = len(tmpList)
        for k in range(0, len_of_sign_list):
            docx_replace(names_and_data[j] + '.docx', list_of_signature[k], tmpList[k])
        j += const_for_add

    j = 0
    for i in range(0, num_of_xlsx):
        tmpList = names_and_data[j + 1: j + const_for_add]  # получаем временный список данных для замены
        num_of_repl_str = len(tmpList)
        for k in range(0, len_of_sign_list):
            excel_replace(names_and_data[j] + '.xlsx', list_of_signature[k], tmpList[k])
        j += const_for_add


    print("End program!\n\n")
    input("Press enter for exit!")
