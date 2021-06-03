from shutil import copy
import yaml
import docx
import openpyxl

def copy_files(num_of_copies_docx:int, example_docx_name, new_docx_name,
              num_of_copies_xlsx:int, example_xlsx_name, new_xlsx_name):
    if not (num_of_copies_docx <= 0):
        for i in range(0, num_of_copies_docx):
            copy(example_docx_name, new_docx_name + str(i)+".docx")
    if not (num_of_copies_xlsx <= 0):
        for i in range(0, num_of_copies_xlsx):
            copy(example_xlsx_name, new_xlsx_name + str(i) + ".xlsx")

def preporation_of_documents(doc_name: str, const1:str, const2:str):
    word_docx = docx.Document(doc_name)
    # поиск сигнатур в обычном тексте
    for paragraph in word_docx.paragraphs:
        # $1 - сигнатура на замену || const1 - вставляемое слово
        paragraph.text = paragraph.text.replace("$a1", const1)
        paragraph.text = paragraph.text.replace("$a2", const2)

        # поиск сигнатур в имеющихся таблицах, актуально к примеру для титульных листов
        for table in word_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace("$a1", const1)
                    cell.text = cell.text.replace("$a2", const2)
    word_docx.save(doc_name)


# функция поиска и замены слов в word docx
def word_replace(doc_name: str, data_list: list):
    word_docx = docx.Document(doc_name)
    list_size = len(data_list)
    i = 0
    while i < list_size:
        # поиск сигнатур в обычном тексте
        table_srch = True
        for paragraph in word_docx.paragraphs:
            # data_list[0] - сигнатура на замену || data_list[1] - вставляемое слово
            paragraph.text = paragraph.text.replace(data_list[i], data_list[i + 1])

        if table_srch:
            # поиск сигнатур в имеющихся таблицах
            for table in word_docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = cell.text.replace(data_list[i], data_list[i + 1])
        i += 2
    word_docx.save(doc_name)

# функция поиска и замены слов в excel xlsx
def excel_replace(doc_name: str, data_list: list):
    excel_doc = openpyxl.load_workbook(doc_name)
    ws = excel_doc["Лист1"] # открытие листа
    list_size = len(data_list)
    i = 0
    while i < list_size:
        #поиск по всем доступным колонкам и стобцам
        for r in range(1, ws.max_row + 1):
            for c in range(1, ws.max_column + 1):
                #получает строку и проверяет ее, если она не пустая пытается выполнить замену
                s = ws.cell(r, c).value
                if s != None and type(s) != int:
                    # data_list[i] - сигнатура на замену || data_list[i+1] - вставляемое слово
                    ws.cell(r, c).value = s.replace(data_list[i], data_list[i+1])

        i += 2
    excel_doc.save(doc_name)


if __name__ == '__main__':
    print("Start program...\n\n")
    # открытие файла yml и считываем данные из него //строки с
    with open("script.yaml", 'r', encoding='utf-8') as stream:
        try:
            load = yaml.safe_load(stream)
            example_wrdfile_name = load['example_wrdfile_name']
            example_exlfile_name = load['example_exlfile_name']
            num_of_docx = load['num_of_docX']
            num_of_xlsx = load['num_of_xlsx']
            new_wrdfile_name = load['new_wrdfile_name']
            new_exlfile_name = load['new_exlfile_name']
            data_list = load['data_list']
        except yaml.YAMLError as exc:
            print(exc)
    print("Finished reading YAML...\n\n")
    # после прочтения файла, создаем копии файлов для заполнения //строки
    print("Start copy fails...\n\n")
    copy_files(num_of_docx, example_wrdfile_name,  new_wrdfile_name,
               num_of_xlsx, example_exlfile_name,  new_exlfile_name)
    print("Сopying files finished...\n\n")

    # подготовка файла, здесь происходит замена сигнатур в файле word, так как изначально в шаблон содержит только
    # сигнатуру $1, $2 и их надо подготовить для заполнения, делается исходя из разработанной концепции
    j:int = 0
    for i in range(0, num_of_docx):
        preporation_of_documents(new_wrdfile_name + str(i) + ".docx", data_list[j], data_list[j + 2])
        j+=4

    print("Start search and past...\n\n")
    # заполняем файлы word
    if not (num_of_docx < 0):
        for i in range(0, num_of_docx):
            word_replace(new_wrdfile_name + str(i) + ".docx", data_list)

    # заполняем файлы docx
    if not (num_of_xlsx < 0):
        for i in range(0, num_of_xlsx):
            excel_replace(new_exlfile_name + str(i) + ".xlsx", data_list)
    print("End program!\n\n")
    input("Press enter for exit!")