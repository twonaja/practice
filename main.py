from shutil import copy
import yaml
import docx
import openpyxl


# функция которая возвращает список индексов у которых в названии
# содержится сигнатура заданая пользователем, входные данные - data_list -> список
# с данными, sign -> сигнатура для поиска
def find_match_indexes(data_list: list, sign: str):
    length_dl = len(data_list)
    index_list = []
    for i in range(0, length_dl):
        if sign in data_list[i]:
            index_list.append(i)
    return index_list


# функция копирования файлов, производит поиск навания файлов с расширением .docx и .xlsx,
# с помощью функции find_match_indexes - возвращает индексы названия новых файлов
# и далее с помощью функции библиотеки shutil -> copy(имя_файла_шаблона, имя_нового_файла)
# производится копирование файлов
def copy_files(example_name: list, dt_list: list, indexes_docx: list, indexes_xlsx):
    # количество документов в списке с расширением .docx
    length_ind_dcx = len(indexes_docx)
    # количество документов в списке с расширением .xlsx
    length_ind_xlx = len(indexes_xlsx)
    for i in range(0, length_ind_dcx):
        # example_name[0] - должно быть название шаблона документа с расширением .docx
        copy(example_name[0], dt_list[indexes_docx[i]]) # indexes_docx[i] - содержит индекс по которому находится в dt_list название нового файла word
    for j in range(0, length_ind_xlx):
        # example_name[1] - должно быть название шаблона документа с расширением .xlsx
        copy(example_name[1], dt_list[indexes_xlsx[j]]) # indexes_xlsx[j]- содержит индекс по которому находится в dt_list название нового файла excel


# функция используется для замены слов и
# зачистки .docx документов - от оставшихся сигнатур
def docx_replace(doc_name: str, signature: str, repl_data: str):
    word_docx = docx.Document(doc_name)
    # поиск сигнатур в обычном тексте
    sign_list_len = len(signature)
    for paragraph in word_docx.paragraphs:
        # signature - сигнатура на замену || repl_data - вставляемая строка
        paragraph.text = paragraph.text.replace(signature, repl_data)
        # поиск сигнатур в имеющихся таблицах, актуально к примеру для титульных листов
        for table in word_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(signature, repl_data)
    word_docx.save(doc_name)


# функция используется для замены слов и
# зачистки .xlsx документов - от оставшихся сигнатур
def excel_replace(doc_name: str, signature: str, repl_data: str):
    excel_doc = openpyxl.load_workbook(doc_name)
    # Лист1 (имя листа в шаблоне не забыть найти более универсальный метод)
    ws = excel_doc["Лист1"]  # открытие листа
    # поиск по всем доступным колонкам и стобцам
    for r in range(1, ws.max_row + 1):
        for c in range(1, ws.max_column + 1):
            # получает строку и проверяет ее, если она не пустая пытается выполнить замену
            s = ws.cell(r, c).value
            # для тестого примера актуально то, что возвращаемая строка не может быть типа - int, если требуется 
            # если требуется менять число, требуется разработать другую функцию, 
            # к примеру if s == 1001: ws.cell(r,c).value = 1000
            if s != None and type(s) != int:
                ws.cell(r, c).value = s.replace(signature, repl_data)
    excel_doc.save(doc_name)


if __name__ == '__main__':
    tmpList = []
    print("Start program...\n\n")
    # часть кода для работы с yaml файлом
    # данные берутся из файла script.yaml - https://github.com/twonaja/practice/blob/main/script.yaml
    with open("script.yaml", 'r', encoding='utf-8') as stream:
        try:
            load = yaml.safe_load(stream)
            # имена шаблонов документов - тип список, [0] - .docx; [1] - .xlsx.
            examples_name = load['examples_name']
            # сигнатуры для вставки
            list_of_signature = load['list_of_signature']
            # имена новых файлов (file_name.docx .... file_name.xlsx) и вставляемые данные - тип список
            names_and_data = load['names_and_data']
        except yaml.YAMLError as exc:
            print(exc)
        # конец работы с yaml файлом

    # indexes_docx - список содержащий индексы имен новых документов с расширением .docx
    ind_docx = find_match_indexes(names_and_data, '.docx')
    # indexes_xlsx - список содержащий индексы имен новых документов с расширением .xlsx
    ind_xlsx = find_match_indexes(names_and_data, '.xlsx')

    # список сигнатур
    # данные списки созданы для того, чтобы знать границы, по какой индекс считывать данные
    ind_beaddoc = find_match_indexes(names_and_data, '0xBEADDOC')
    ind_beadxl = find_match_indexes(names_and_data, '0xBEADXL')
    # копируем файлы
    copy_files(examples_name, names_and_data, ind_docx, ind_xlsx)

    # заменяем сигнатуры
    # индекс для замены, 
    ind_docx_len = len(ind_docx)
    ind_xlsx_len = len(ind_xlsx)
    for i in range(0, ind_docx_len):
        tmpList = names_and_data[ind_docx[i] + 1: ind_beaddoc[i]] # получаем временный список данных для замены
        num_of_repl_str = len(tmpList)
        for j in range(0, num_of_repl_str):
            # заменяем сигнатуру на нужную строку
            docx_replace(names_and_data[ind_docx[i]], list_of_signature[j], tmpList[j])

    for i in range(0, ind_xlsx_len):
        tmpList = names_and_data[ind_xlsx[i] + 1: ind_beadxl[i]] # получаем временный список данных для замены
        num_of_repl_str = len(tmpList)
        for j in range(0, num_of_repl_str):
            # заменяем сигнатуру на нужную строку
            excel_replace(names_and_data[ind_xlsx[i]], list_of_signature[j], tmpList[j])

    # отчистка файлов от оставшихся сигнатур
    for i in range(0, ind_docx_len):
        len_of_lst_sign = len(list_of_signature)
        for j in range(0, len_of_lst_sign):
            # заменяем сигнатуру на нужную строку
            docx_replace(names_and_data[ind_docx[i]], list_of_signature[j], ' ')
    # отчистка файлов от оставшихся сигнатур
    for i in range(0, ind_xlsx_len):
        len_of_lst_sign = len(list_of_signature)
        for j in range(0, len_of_lst_sign):
            # заменяем сигнатуру на нужную строку
            excel_replace(names_and_data[ind_xlsx[i]], list_of_signature[j], ' ')

    print("End program!\n\n")
    input("Press enter for exit!")
